import os
import urllib.request
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def build_srs_document():
    doc = Document()

    # Define color scheme
    COLOR_PRIMARY = RGBColor(26, 82, 118)   # Deep Navy Blue
    COLOR_SECONDARY = RGBColor(41, 128, 185) # Soft Blue
    COLOR_ORANGE = RGBColor(211, 84, 0)     # LPU Orange
    COLOR_TEXT = RGBColor(44, 62, 80)       # Dark Slate
    COLOR_LIGHT_GREY = RGBColor(127, 140, 141)

    # Set page margins (1 inch on all sides)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # --- COVER PAGE ---
    # Try to download LPU logo
    logo_path = "lpu_logo.png"
    if not os.path.exists(logo_path):
        try:
            logo_url = "https://upload.wikimedia.org/wikipedia/commons/thumb/a/a4/Lovely_Professional_University_logo.svg/1200px-Lovely_Professional_University_logo.svg.png"
            urllib.request.urlretrieve(logo_url, logo_path)
        except Exception:
            logo_path = None

    # LPU Logo
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(10)
    if logo_path and os.path.exists(logo_path):
        p_logo.add_run().add_picture(logo_path, width=Inches(1.8))

    # University Title
    p_univ = doc.add_paragraph()
    p_univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_univ.paragraph_format.space_before = Pt(12)
    run_univ = p_univ.add_run("LOVELY PROFESSIONAL UNIVERSITY\n")
    run_univ.font.name = 'Arial'
    run_univ.font.size = Pt(18)
    run_univ.font.bold = True
    run_univ.font.color.rgb = COLOR_ORANGE

    run_dept = p_univ.add_run("Department of Computer Science & Engineering")
    run_dept.font.name = 'Arial'
    run_dept.font.size = Pt(12)
    run_dept.font.bold = True
    run_dept.font.color.rgb = COLOR_LIGHT_GREY

    # Divider
    p_div = doc.add_paragraph()
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_div.add_run("____________________________________________________").font.color.rgb = COLOR_LIGHT_GREY

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(24)
    run_title = p_title.add_run("SOFTWARE REQUIREMENTS SPECIFICATION\n")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY

    run_sub = p_title.add_run("🅿️ PARK-KARO: Smart Parking Management System")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(16)
    run_sub.font.bold = True
    run_sub.font.color.rgb = COLOR_SECONDARY

    # Divider
    p_div2 = doc.add_paragraph()
    p_div2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_div2.add_run("____________________________________________________").font.color.rgb = COLOR_LIGHT_GREY

    # Student Info
    p_student = doc.add_paragraph()
    p_student.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_student.paragraph_format.space_before = Pt(36)
    run_student_label = p_student.add_run("Submitted By:\n")
    run_student_label.font.name = 'Arial'
    run_student_label.font.size = Pt(12)
    run_student_label.font.italic = True
    run_student_label.font.color.rgb = COLOR_TEXT

    run_student_name = p_student.add_run("SUMIT KUMAR\n")
    run_student_name.font.name = 'Arial'
    run_student_name.font.size = Pt(16)
    run_student_name.font.bold = True
    run_student_name.font.color.rgb = COLOR_PRIMARY

    run_course = p_student.add_run("Bachelor of Technology (Computer Science & Engineering)\n")
    run_course.font.name = 'Arial'
    run_course.font.size = Pt(12)
    run_course.font.color.rgb = COLOR_TEXT

    run_guide = p_student.add_run("Under the Guidance of Lovely Professional University, Punjab")
    run_guide.font.name = 'Arial'
    run_guide.font.size = Pt(11)
    run_guide.font.italic = True
    run_guide.font.color.rgb = COLOR_LIGHT_GREY

    # Year
    p_year = doc.add_paragraph()
    p_year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_year.paragraph_format.space_before = Pt(48)
    run_year = p_year.add_run("Academic Year: 2026")
    run_year.font.name = 'Arial'
    run_year.font.size = Pt(11)
    run_year.font.bold = True
    run_year.font.color.rgb = COLOR_TEXT

    # --- PAGE BREAK TO TABLE OF CONTENTS ---
    doc.add_page_break()

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY
        return p

    def add_paragraph(text, bold_prefix="", italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.font.name = 'Arial'
            run_b.font.size = Pt(11)
            run_b.font.bold = True
            run_b.font.color.rgb = COLOR_TEXT
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.italic = italic
        run.font.color.rgb = COLOR_TEXT
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.font.name = 'Arial'
            run_b.font.size = Pt(11)
            run_b.font.bold = True
            run_b.font.color.rgb = COLOR_TEXT
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_TEXT
        return p

    # Table of Contents
    add_heading_1("📋 Table of Contents")
    add_paragraph("1. Problem Statement", bold_prefix="Section 1: ")
    add_paragraph("2. The Solution (ParkKaro)", bold_prefix="Section 2: ")
    add_paragraph("3. Key Features", bold_prefix="Section 3: ")
    add_paragraph("4. High-Level Design (HLD)", bold_prefix="Section 4: ")
    add_paragraph("5. Low-Level Design (LLD)", bold_prefix="Section 5: ")
    add_paragraph("6. Database & ERD Design", bold_prefix="Section 6: ")
    add_paragraph("7. Technology Stack", bold_prefix="Section 7: ")
    add_paragraph("8. Output Screens & User Interface Flow", bold_prefix="Section 8: ")
    add_paragraph("9. Future Scope", bold_prefix="Section 9: ")
    add_paragraph("10. Conclusion", bold_prefix="Section 10: ")

    doc.add_page_break()

    # Section 1
    add_heading_1("🔍 1. Problem Statement")
    add_paragraph("In rapidly growing cities, hunting for a parking spot is one of the most frustrating daily hassles for drivers. Whether visiting a shopping mall, a college, a hospital, or an office complex, finding a space to park is painful and inefficient.")
    add_heading_2("⚠️ The Key Pain Points:")
    add_bullet("Drivers circle around parking lots repeatedly, wasting an average of 15–20 minutes and burning fuel unnecessarily. This also increases traffic congestion outside venues.", bold_prefix="Wasted Time & Fuel: ")
    add_bullet("Drivers have no way of knowing if a parking garage is already full until they drive all the way inside.", bold_prefix="Zero Visibility (Blind Spots): ")
    add_bullet("Old-school paper tickets and cash-based entries are slow, prone to human error, and lack safety verification.", bold_prefix="Outdated Manual Tracking: ")
    add_bullet("If plans change, drivers cannot easily cancel their booking or extend their parking duration remotely.", bold_prefix="No Flexibility: ")

    # Section 2
    add_heading_1("💡 2. The Solution (ParkKaro)")
    add_paragraph("ParkKaro is a smart, web-based, real-time parking slot booking system designed to make parking as easy as booking a movie ticket! It brings transparency, speed, and modern convenience to drivers and parking operators alike.")
    
    add_heading_2("✅ How ParkKaro Solves the Problem:")
    add_bullet("Drivers can view real-time availability of slots before arriving.", bold_prefix="Live Spot Maps: ")
    add_bullet("Users select their preferred parking spot using an intuitive visual map.", bold_prefix="Instant Booking: ")
    add_bullet("Need more time? Extend your booking with one click. Change of plans? Cancel instantly with a fair refund.", bold_prefix="Hassle-Free Extensions & Cancellations: ")
    add_bullet("Every booking generates a unique digital QR Code that can be scanned for quick, paperless check-in and check-out.", bold_prefix="Contactless QR Tickets: ")

    # Section 3
    add_heading_1("✨ 3. Key Features")
    add_paragraph("Our system is loaded with features categorized for two primary groups: Drivers (Users) and Parking Operators (Admins).")
    
    add_heading_2("🚗 User Features:")
    add_bullet("A visual, color-coded grid map showing available slots (Green), booked slots (Red), and selected slots (Blue) with smooth hover effects.", bold_prefix="Interactive Parking Map: ")
    add_bullet("Search and filter parking spaces by City, Area, or Location Type (Mall, Hotel, College, Office, etc.).", bold_prefix="Smart Search: ")
    add_bullet("Extend an active parking reservation directly from the dashboard if you are running late.", bold_prefix="Dynamic Time Extensions: ")
    add_bullet("Cancel an active reservation up to the start time with an automatic refund (minus a standard 10% fee).", bold_prefix="Fair Cancellation Policy: ")
    add_bullet("Instant ticket generation with a secure, scan-ready QR Code containing booking details.", bold_prefix="QR E-Ticket: ")

    add_heading_2("🛡️ Admin & Operator Features:")
    add_bullet("A dedicated page for operators to scan or enter ticket codes to instantly verify active status and mark check-ins/check-outs.", bold_prefix="QR Code Scanner / Verifier: ")
    add_bullet("View real-time occupancy statistics across all locations from a single dashboard.", bold_prefix="Live Slot Monitoring: ")

    # Section 4
    add_heading_1("🏗️ 4. High-Level Design (HLD)")
    add_paragraph("The High-Level Design explains how the different pieces of ParkKaro connect and communicate with each other. The application is built using the robust Model-View-Template (MVT) architecture.")
    add_paragraph("The system consists of three major blocks: Client Browser, Django Web Server Core (comprising Views, Templates, and Models), and the Database. When a user requests a page, Django Views execute business logic, communicate with the Database using Django ORM, and render the final template to the client browser while integrating the Python-QRcode API to deliver dynamic verification tickets.")

    # Section 5
    add_heading_1("⚙️ 5. Low-Level Design (LLD)")
    add_paragraph("The Low-Level Design goes deep into the coding logic, rules, and mathematical calculations that make ParkKaro work reliably.")
    
    add_heading_2("🧮 Crucial System Algorithms:")
    add_paragraph("Total Amount = Total Hours x Hourly Rate (e.g. Rs 20). The hours are calculated dynamically as the difference between the starting and ending times.", bold_prefix="1. Dynamic Price Calculator: ")
    add_paragraph("If a user cancels a booking, they get a refund calculated automatically: Refund = Paid Amount - max(5.0, Paid Amount x 0.10). A minimum fee of Rs 5 is retained to cover transaction processing.", bold_prefix="2. Smart Refund Calculation (10% Fee): ")
    add_paragraph("The system checks if the selected slot is free for the extended hours. If free, it adds the extra hours to end_time, calculates the extra fee, and updates the existing ticket.", bold_prefix="3. Time Extension Logic: ")

    # Section 6
    add_heading_1("🗄️ 6. Database & ERD Design")
    add_paragraph("Our database uses clean relations to store and link users, locations, parking slots, and reservations.")
    
    add_heading_2("📋 Database Schema Fields:")
    add_bullet("id (PK), username, email, password", bold_prefix="User Table: ")
    add_bullet("id (PK), name, city, area, location_type", bold_prefix="Location Table: ")
    add_bullet("id (PK), location_id (FK), slot_number, is_available", bold_prefix="ParkingSlot Table: ")
    add_bullet("id (PK), user_id (FK), slot_id (FK), vehicle_number, booking_date, start_time, end_time, total_hours, total_amount, status", bold_prefix="Booking Table: ")

    # Create Table
    add_heading_2("📋 Detailed Schema Table")
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Table Name'
    hdr_cells[1].text = 'Attribute (Field)'
    hdr_cells[2].text = 'Data Type'
    hdr_cells[3].text = 'Description'
    
    # Set headers bold
    for cell in hdr_cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.name = 'Arial'

    data = [
        ('Location', 'id', 'Integer (PK)', 'Unique identifier for parking venue'),
        ('Location', 'name', 'Varchar', 'Name of the venue (e.g. LPU Block 34)'),
        ('Location', 'city', 'Varchar', 'City name (e.g. Jalandhar)'),
        ('ParkingSlot', 'id', 'Integer (PK)', 'Unique identifier for slot'),
        ('ParkingSlot', 'slot_number', 'Varchar', 'Human-readable spot label (e.g. S-4)'),
        ('ParkingSlot', 'is_available', 'Boolean', 'Availability status (True/False)'),
        ('Booking', 'id', 'Integer (PK)', 'Unique identifier for booking'),
        ('Booking', 'vehicle_number', 'Varchar', 'Car/Bike plate number'),
        ('Booking', 'total_amount', 'Float', 'Price of reservation'),
        ('Booking', 'status', 'Varchar', 'PENDING / ACTIVE / CANCELLED')
    ]

    for item in data:
        row_cells = table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = item[2]
        row_cells[3].text = item[3]
        for cell in row_cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9.5)

    # Section 7
    add_heading_1("🛠️ 7. Technology Stack")
    add_bullet("HTML5, CSS3 (Glassmorphic design), Vanilla JS (instant state rendering)", bold_prefix="Frontend: ")
    add_bullet("Python 3 & Django Framework (built-in security, ORM database handling)", bold_prefix="Backend: ")
    add_bullet("SQLite3 (Development) / PostgreSQL (Production)", bold_prefix="Database: ")
    add_bullet("Python-QRcode Library", bold_prefix="APIs & Helpers: ")

    # Section 8
    add_heading_1("🖥️ 8. Output Screens & User Interface Flow")
    add_paragraph("Each screen is fully responsive with a high-end glassmorphic UI.")
    add_bullet("Sleek authentication forms.", bold_prefix="1. Login & Registration: ")
    add_bullet("Interactive search cards by city or venue type.", bold_prefix="2. User Dashboard: ")
    add_bullet("Live physical layout mapping with color indicators (Green/Red/Blue).", bold_prefix="3. Interactive Slot Map: ")
    add_bullet("Summarizes slots, displays checkout totals, and handles mocks payments.", bold_prefix="4. Booking & Payment checkout: ")
    add_bullet("Enables cancellations, duration extensions, and displays QR Tickets.", bold_prefix="5. Booking History & Ticket Hub: ")
    add_bullet("Special operator screen to scan and verify QR codes for entry/exit.", bold_prefix="6. Ticket Verification: ")

    # Section 9
    add_heading_1("🚀 9. Future Scope")
    add_bullet("Connect ultrasonic hardware sensors with ESP32 to detect presence automatically.", bold_prefix="1. IoT Sensors: ")
    add_bullet("Automated License Plate Recognition using camera-based AI.", bold_prefix="2. ALPR AI Cameras: ")
    add_bullet("Let users reserve slots with EV charging stations.", bold_prefix="3. EV Smart Charging: ")
    add_bullet("Surge-pricing during peak traffic hours.", bold_prefix="4. Dynamic Pricing: ")

    # Section 10
    add_heading_1("🎯 10. Conclusion")
    add_paragraph("ParkKaro successfully transforms the tedious, paper-heavy chore of searching for parking into a fast, transparent, and completely digital experience. By giving drivers real-time visual power and contactless QR check-ins, the system reduces congestion, saves fuel, and delivers absolute financial transparency for operators.")

    # Footer Centered
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.space_before = Pt(36)
    run_f = p_foot.add_run("© 2026 ParkKaro by Sumit Kumar | Lovely Professional University")
    run_f.font.name = 'Arial'
    run_f.font.bold = True
    run_f.font.size = Pt(10)
    run_f.font.color.rgb = COLOR_PRIMARY

    output_filename = "SRS_Document.docx"
    doc.save(output_filename)
    print(f"Document successfully created as {output_filename}!")

if __name__ == "__main__":
    build_srs_document()
